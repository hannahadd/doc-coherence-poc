"""
Extraction de texte unifiée pour tous les formats supportés.

Formats :
  - PDF          : pypdf  (natif, sans OCR)
  - Image        : Tesseract OCR via pytesseract (JPG, PNG, TIFF, BMP, WEBP)
  - DOCX         : python-docx
  - TXT          : lecture directe UTF-8

Usage :
    from src.doc_reader import extract_from_bytes, ExtractedDoc
    doc = extract_from_bytes(file_bytes, "mon_cv.jpg")
    print(doc.text, doc.source_type, doc.ocr_used)
"""

from __future__ import annotations

import io
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

from src.pdf_utils import extract_text_from_pdf_bytes, clean_text


@dataclass
class ExtractedDoc:
    text: str
    source_type: str          # "pdf" | "image" | "docx" | "txt"
    num_pages: Optional[int] = None
    ocr_used: bool = False
    warning: Optional[str] = None


# ── Extensions reconnues ──────────────────────────────────────────────────────

IMAGE_EXTS = {".jpg", ".jpeg", ".png", ".tiff", ".tif", ".bmp", ".webp"}
DOCX_EXTS  = {".docx"}
PDF_EXTS   = {".pdf"}
TXT_EXTS   = {".txt"}

ALL_SUPPORTED = PDF_EXTS | IMAGE_EXTS | DOCX_EXTS | TXT_EXTS


# ── Extracteurs internes ──────────────────────────────────────────────────────

def _extract_pdf(file_bytes: bytes) -> ExtractedDoc:
    result = extract_text_from_pdf_bytes(file_bytes)
    return ExtractedDoc(
        text=result.text,
        source_type="pdf",
        num_pages=result.num_pages,
        ocr_used=False,
    )


def _extract_image(file_bytes: bytes) -> ExtractedDoc:
    """OCR via Tesseract (doit être installé sur le système : tesseract-ocr + tesseract-ocr-fra)."""
    try:
        import pytesseract
        from PIL import Image
    except ImportError as e:
        return ExtractedDoc(
            text="",
            source_type="image",
            ocr_used=True,
            warning=f"Dépendance manquante pour l'OCR : {e}. Installe pytesseract et Pillow.",
        )

    try:
        img = Image.open(io.BytesIO(file_bytes))
        # lang="fra+eng" : français en priorité, anglais en fallback
        raw = pytesseract.image_to_string(img, lang="fra+eng")
    except pytesseract.TesseractNotFoundError:
        return ExtractedDoc(
            text="",
            source_type="image",
            ocr_used=True,
            warning=(
                "Tesseract n'est pas installé sur ce système. "
                "Sur macOS : brew install tesseract tesseract-lang. "
                "Sur Linux/Docker : apt-get install tesseract-ocr tesseract-ocr-fra."
            ),
        )
    except Exception as e:
        return ExtractedDoc(
            text="",
            source_type="image",
            ocr_used=True,
            warning=f"Erreur OCR : {e}",
        )

    return ExtractedDoc(
        text=clean_text(raw),
        source_type="image",
        ocr_used=True,
    )


def _extract_docx(file_bytes: bytes) -> ExtractedDoc:
    try:
        from docx import Document
    except ImportError as e:
        return ExtractedDoc(
            text="",
            source_type="docx",
            warning=f"Dépendance manquante : {e}. Installe python-docx.",
        )

    try:
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        # Inclut aussi le texte dans les tableaux
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    t = cell.text.strip()
                    if t:
                        paragraphs.append(t)
        raw = "\n".join(paragraphs)
    except Exception as e:
        return ExtractedDoc(
            text="",
            source_type="docx",
            warning=f"Erreur lecture DOCX : {e}",
        )

    return ExtractedDoc(
        text=clean_text(raw),
        source_type="docx",
    )


def _extract_txt(file_bytes: bytes) -> ExtractedDoc:
    try:
        raw = file_bytes.decode("utf-8", errors="replace")
    except Exception as e:
        return ExtractedDoc(text="", source_type="txt", warning=f"Erreur lecture TXT : {e}")
    return ExtractedDoc(text=clean_text(raw), source_type="txt")


# ── Point d'entrée public ─────────────────────────────────────────────────────

def extract_from_bytes(file_bytes: bytes, filename: str) -> ExtractedDoc:
    """
    Détecte le format via l'extension du nom de fichier et extrait le texte.

    Args:
        file_bytes : contenu brut du fichier
        filename   : nom original du fichier (utilisé pour détecter le format)

    Returns:
        ExtractedDoc avec .text, .source_type, .ocr_used, .warning
    """
    ext = Path(filename).suffix.lower()

    if ext in PDF_EXTS:
        return _extract_pdf(file_bytes)
    elif ext in IMAGE_EXTS:
        return _extract_image(file_bytes)
    elif ext in DOCX_EXTS:
        return _extract_docx(file_bytes)
    elif ext in TXT_EXTS:
        return _extract_txt(file_bytes)
    else:
        return ExtractedDoc(
            text="",
            source_type="unknown",
            warning=f"Format non supporté : '{ext}'. Formats acceptés : PDF, image (JPG/PNG/TIFF/BMP/WEBP), DOCX, TXT.",
        )
